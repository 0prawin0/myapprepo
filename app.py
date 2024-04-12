import openai
import streamlit as st
from docx import Document
from google.cloud import storage
import os
import uuid
import pdfkit
from docx.shared import Pt

openai_api_key = os.getenv('OPENAI_API_KEY')
if openai_api_key is None:
    raise ValueError("OPENAI_API_KEY environment variable is not set.")

st.title("🧑‍💻 Job Genius JD Creator")
st.write("""
Hi, I am your JD Creator🤖 
Let me help you create a job description that would suit your requirements.
""")

# Initialize the session state keys
if "messages" not in st.session_state:
    st.session_state["messages"] = [
        {"role": "system", "content": "The following is a conversation with an AI assistant helping to create a job description. The assistant collects the following data one by one by interacting with the user by asking questions to create the best job description: company_name, job_title, key_skills, soft_skills, location, desired_experience, preferred_experience, about_the_team ."}
    ]

if "job_title_asked" not in st.session_state:
    st.session_state["job_title_asked"] = False

# Ask for the job title if not asked before
if not st.session_state["job_title_asked"]:
    st.session_state["messages"].append({"role": "assistant", "content": "What is the job title?"})
    st.session_state["job_title_asked"] = True

# Handle user input
if prompt := st.chat_input():
    openai.api_key = openai_api_key
    st.session_state["messages"].append({"role": "user", "content": prompt})
    response = openai.ChatCompletion.create(model="gpt-3.5-turbo-0613", messages=st.session_state["messages"])
    msg = response.choices[0].message
    st.session_state["messages"].append(msg)
    st.chat_message("assistant").write(msg["content"])

# Function to save JD to a PDF file and upload to Google Cloud Storage
def save_jd_to_bucket(jd_text, bucket_name, job_role, designation, company_name, location):
    # Create a unique file name
    unique_id = str(uuid.uuid4())
    word_file_name = f"{job_role}-{designation}-{unique_id}.docx"
    pdf_file_name = f"{job_role}-{designation}-{unique_id}.pdf"

    # Create a Word document
    doc = Document()
    
    # Add the entire conversation
    for msg in st.session_state["messages"]:
        line = msg['content']
        doc.add_paragraph(line)

    # Save the Word document
    doc.save(word_file_name)

    # Create a new Word document for the PDF conversion
    doc_pdf = Document()
    
    # Add company name, job role, and location in bold and larger font
    company_paragraph = doc_pdf.add_paragraph()
    company_run = company_paragraph.add_run(f"Company Name: {company_name}\nJob Title: {job_role}\nLocation: {location}\n")
    company_run.bold = True
    company_run.font.size = Pt(16)  # Change the font size as needed

    doc_pdf.add_paragraph()  # Add an empty line for spacing

    # Add only the final job description content for the PDF
    doc_pdf.add_paragraph(jd_text)

    # Save the PDF-specific Word document
    doc_pdf.save(pdf_file_name)

    # Convert the PDF-specific Word document to PDF
    pdfkit.from_file(pdf_file_name, pdf_file_name)

    # Upload the PDF to Google Cloud Storage
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(pdf_file_name)
    blob.upload_from_filename(pdf_file_name)

    # Remove the files after uploading
    os.remove(word_file_name)
    os.remove(pdf_file_name)

    return f"Job description saved to {bucket_name}/{pdf_file_name}"

# Button to submit the job description
if st.button('Submit'):
    jd_text = '\n'.join([msg['content'] for msg in st.session_state["messages"] if msg['role'] == 'assistant' and 'final job description' in msg['content'].lower()])
    
    # Extract job role, designation, company name, and location from the job description
    job_role = "Unknown"  # Default value
    designation = "Unknown"  # Default value
    company_name = "Unknown"  # Default value
    location = "Unknown"  # Default value
    for msg in st.session_state["messages"]:
        if msg['role'] == 'assistant':
            content = msg['content']
            if content.startswith("What is the job title?") and ":" in content:
                job_role = content.split(":")[1].strip()
            elif content.startswith("What is the designation?") and ":" in content:
                designation = content.split(":")[1].strip()
            elif content.startswith("What is the company name?") and ":" in content:
                company_name = content.split(":")[1].strip()
            elif content.startswith("Where is the location?") and ":" in content:
                location = content.split(":")[1].strip()

    result = save_jd_to_bucket(jd_text, 'jd_storage_bucket', job_role, designation, company_name, location)
    st.success(result)
